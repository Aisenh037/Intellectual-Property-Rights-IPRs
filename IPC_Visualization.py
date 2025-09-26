# Enhanced version with research paper-ready features
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.io as pio
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import warnings
import os
from pathlib import Path
import networkx as nx
import requests
from additional_visualizations import create_india_applications_choropleth, create_india_grants_choropleth

# Research paper styling
plt.style.use('seaborn-v0_8-whitegrid')
sns.set_palette("husl")
warnings.filterwarnings('ignore')

# Enhanced configuration for research output
class ResearchConfig:
    def __init__(self):
        self.DATA_PATHS = {
            'applications': r"C:\Users\ASUS\Desktop\Research Work\Main_IPC_Working Dir\SONAL_FINAL_IPC_FILES\Green_AIN.csv",
            'grants': r"C:\Users\ASUS\Desktop\Research Work\Main_IPC_Working Dir\SONAL_FINAL_IPC_FILES\Green_GIN.csv"
        }
        
        # Research paper color scheme
        self.COLOR_SCHEME = {
            'primary': '#2E8B57',      # Green (main theme)
            'secondary': '#1F77B4',    # Blue
            'tertiary': '#FF7F0E',     # Orange
            'accent': '#D62728',       # Red
            'neutral': '#9467BD'       # Purple
        }
        
        # Output directories
        self.OUTPUT_DIR = "research_output"
        self.FIGURES_DIR = f"{self.OUTPUT_DIR}/figures"
        self.DATA_DIR = f"{self.OUTPUT_DIR}/data"
        self.HTML_DIR = f"{self.OUTPUT_DIR}/html_reports"
        
        # Create directories
        self._create_directories()
    
    def _create_directories(self):
        """Create output directories for research materials"""
        for directory in [self.OUTPUT_DIR, self.FIGURES_DIR, self.DATA_DIR, self.HTML_DIR]:
            Path(directory).mkdir(parents=True, exist_ok=True)

config = ResearchConfig()

# Enhanced data preprocessing with research-focused features
class ResearchPatentPreprocessor:
    def __init__(self, app_df, grant_df):
        self.app_df = app_df
        self.grant_df = grant_df
        self.ipc_sections = {
            'A': 'Human Necessities',
            'B': 'Performing Operations; Transporting',
            'C': 'Chemistry; Metallurgy',
            'D': 'Textiles; Paper',
            'E': 'Fixed Constructions',
            'F': 'Mechanical Engineering; Lighting; Heating; Weapons; Blasting',
            'G': 'Physics',
            'H': 'Electricity'
        }
    
    def enhanced_ipc_analysis(self, df, column_name='Cleaned_ICR'):
        """Comprehensive IPC analysis with research insights"""
        # Clean IPC column if not already cleaned
        if 'IPC Revised (ICR)' in df.columns and column_name not in df.columns:
            df[column_name] = df['IPC Revised (ICR)'].astype(str).str.strip()
            df[column_name] = df[column_name].str.split(';').str[0]  # Take first IPC code if multiple
            df[column_name] = df[column_name].str.replace(r'\s+', '', regex=True)  # Remove extra spaces
        
        # Extract IPC hierarchy
        df[[f'{column_name}_Section', f'{column_name}_Class', 
            f'{column_name}_Subclass', f'{column_name}_Group']] = self.extract_ipc_hierarchy(df[column_name])
        
        # Add descriptive section names
        df[f'{column_name}_Section_Name'] = df[f'{column_name}_Section'].map(self.ipc_sections)
        
        # Clean additional columns for advanced visualizations
        tech_cols = ['Tech Domain', 'Tech Sub Domain', 'Industry']
        for col in tech_cols:
            if col in df.columns:
                df[col] = df[col].astype(str).str.strip().str.title()  # Capitalize first letter
                df[col] = df[col].replace('Nan', 'Unknown').replace('None', 'Unknown')
        
        # Standardize Assignee Country
        if 'Assignee Country' in df.columns:
            country_mapping = {
                'USA': 'US', 'United States': 'US', 'United States Of America': 'US',
                'India': 'IN', 'Uae': 'AE', 'China': 'CN', 'Japan': 'JP',
                'Germany': 'DE', 'United Kingdom': 'GB', 'France': 'FR',
                'Canada': 'CA', 'Australia': 'AU', 'South Korea': 'KR'
            }
            df['Assignee Country'] = df['Assignee Country'].astype(str).str.strip().str.upper()
            df['Assignee Country'] = df['Assignee Country'].map(country_mapping).fillna(df['Assignee Country'])
            df['Assignee Country'] = df['Assignee Country'].replace('NAN', 'Unknown').replace('NONE', 'Unknown')
        
        # Clean Region if present
        if 'Region' in df.columns:
            df['Region'] = df['Region'].astype(str).str.strip().str.title()
            df['Region'] = df['Region'].replace('Nan', 'Unknown').replace('None', 'Unknown')
        
        return df

    def extract_ipc_hierarchy(self, ipc_series):
        """Enhanced IPC extraction with validation"""
        def extract_levels(ipc_code):
            if pd.isna(ipc_code):
                return pd.Series([None, None, None, None])

            codes = str(ipc_code).split(';')
            if codes:
                main_code = codes[0].strip()
                # Validate IPC format: should start with valid section letter followed by 2 digits minimum
                if len(main_code) < 3 or not main_code[0].isalpha() or not main_code[1:3].isdigit():
                    return pd.Series([None, None, None, None])

                section = main_code[0].upper()
                # Check if section is valid (A-H)
                if section not in self.ipc_sections:
                    return pd.Series([None, None, None, None])

                klass = main_code[0:3]  # Always valid if we pass the check above

                # Subclass should be section + 3 digits
                if len(main_code) >= 4 and main_code[3].isdigit():
                    subclass = main_code[0:4]
                else:
                    subclass = None

                # Group is everything before '/'
                group = main_code.split('/')[0] if '/' in main_code else main_code

                return pd.Series([section, klass, subclass, group])
            return pd.Series([None, None, None, None])

        return ipc_series.apply(extract_levels)
    
    def calculate_research_metrics(self):
        """Calculate metrics suitable for research publication"""
        metrics = {}
        
        # Basic statistics
        metrics['total_applications'] = len(self.app_df)
        metrics['total_grants'] = len(self.grant_df)
        metrics['overall_grant_rate'] = (metrics['total_grants'] / metrics['total_applications'] * 100) if metrics['total_applications'] > 0 else 0
        
        # Temporal analysis
        valid_years = self.app_df['Application_Year'].dropna()
        if len(valid_years) > 0:
            metrics['time_span'] = {
                'start': valid_years.min(),
                'end': valid_years.max(),
                'duration': valid_years.max() - valid_years.min()
            }
        else:
            metrics['time_span'] = {'start': None, 'end': None, 'duration': 0}
        
        # Technology concentration
        if 'Cleaned_ICR_Section' in self.app_df.columns:
            top_tech_apps = self.app_df['Cleaned_ICR_Section'].value_counts().head(3)
            metrics['technology_concentration'] = {
                'top_technologies': top_tech_apps.to_dict(),
                'hhi': self.calculate_hhi(self.app_df['Cleaned_ICR_Section'])
            }
        else:
            metrics['technology_concentration'] = {'top_technologies': {}, 'hhi': 0}
        
        return metrics
    
    def calculate_hhi(self, series):
        """Calculate Herfindahl-Hirschman Index for concentration"""
        value_counts = series.value_counts(normalize=True)
        return (value_counts ** 2).sum() * 10000

    def process_geography_and_industry(self):
        """Aggregate geography and industry data with grant rates"""
        # Merge app and grant on Application No. for grant status
        merge_key = 'Application No.'
        if merge_key not in self.app_df.columns or merge_key not in self.grant_df.columns:
            print(f"Warning: {merge_key} not found. Skipping grant rate calculations.")
            return {}

        merged = pd.merge(
            self.app_df[[merge_key, 'Assignee Country', 'Region', 'Industry', 'Tech Domain', 'Tech Sub Domain']],
            self.grant_df[[merge_key]],
            on=merge_key,
            how='left',
            indicator=True
        )
        merged['Granted'] = (merged['_merge'] == 'both').astype(int)

        aggregations = {}

        # Geography: by Assignee Country
        if 'Assignee Country' in merged.columns:
            country_agg = merged.groupby('Assignee Country').agg({
                'Granted': ['count', 'sum'],
                'Application No.': 'nunique'  # Unique apps
            }).round(2)
            country_agg.columns = ['Total_Apps', 'Grants', 'Unique_Apps']
            country_agg['Grant_Rate'] = (country_agg['Grants'] / country_agg['Total_Apps'] * 100).round(2)
            country_agg = country_agg.sort_values('Total_Apps', ascending=False)
            aggregations['country_summary'] = country_agg

        # Geography: by Region
        if 'Region' in merged.columns:
            region_agg = merged.groupby('Region').agg({
                'Granted': ['count', 'sum']
            }).round(2)
            region_agg.columns = ['Total_Apps', 'Grants']
            region_agg['Grant_Rate'] = (region_agg['Grants'] / region_agg['Total_Apps'] * 100).round(2)
            region_agg = region_agg.sort_values('Total_Apps', ascending=False)
            aggregations['region_summary'] = region_agg

        # Industry
        if 'Industry' in merged.columns:
            industry_agg = merged.groupby('Industry').agg({
                'Granted': ['count', 'sum']
            }).round(2)
            industry_agg.columns = ['Total_Apps', 'Grants']
            industry_agg['Grant_Rate'] = (industry_agg['Grants'] / industry_agg['Total_Apps'] * 100).round(2)
            industry_agg = industry_agg.sort_values('Total_Apps', ascending=False)
            aggregations['industry_summary'] = industry_agg

        # Tech Domain
        if 'Tech Domain' in merged.columns:
            tech_domain_agg = merged.groupby('Tech Domain').agg({
                'Granted': ['count', 'sum']
            }).round(2)
            tech_domain_agg.columns = ['Total_Apps', 'Grants']
            tech_domain_agg['Grant_Rate'] = (tech_domain_agg['Grants'] / tech_domain_agg['Total_Apps'] * 100).round(2)
            tech_domain_agg = tech_domain_agg.sort_values('Total_Apps', ascending=False)
            aggregations['tech_domain_summary'] = tech_domain_agg

        # Tech Sub Domain (top 20 to avoid too many)
        if 'Tech Sub Domain' in merged.columns:
            top_subdomains = merged['Tech Sub Domain'].value_counts().head(20).index
            sub_domain_agg = merged[merged['Tech Sub Domain'].isin(top_subdomains)].groupby('Tech Sub Domain').agg({
                'Granted': ['count', 'sum']
            }).round(2)
            sub_domain_agg.columns = ['Total_Apps', 'Grants']
            sub_domain_agg['Grant_Rate'] = (sub_domain_agg['Grants'] / sub_domain_agg['Total_Apps'] * 100).round(2)
            sub_domain_agg = sub_domain_agg.sort_values('Total_Apps', ascending=False)
            aggregations['tech_subdomain_summary'] = sub_domain_agg

        return aggregations

# Advanced research visualizer
class ResearchVisualizer:
    def __init__(self, config):
        self.config = config
        self.ipc_sections = {
            'A': 'Human Necessities',
            'B': 'Performing Operations; Transporting',
            'C': 'Chemistry; Metallurgy',
            'D': 'Textiles; Paper',
            'E': 'Fixed Constructions',
            'F': 'Mechanical Engineering; Lighting; Heating; Weapons; Blasting',
            'G': 'Physics',
            'H': 'Electricity'
        }
        self.setup_research_style()
    
    def setup_research_style(self):
        """Setup publication-quality styling"""
        # Plotly template for research papers
        pio.templates["research"] = go.layout.Template(
            layout=go.Layout(
                font=dict(family="Times New Roman", size=12),
                title=dict(font=dict(size=16, family="Times New Roman", color='black')),
                xaxis=dict(title=dict(font=dict(size=14)), tickfont=dict(size=12)),
                yaxis=dict(title=dict(font=dict(size=14)), tickfont=dict(size=12)),
                legend=dict(font=dict(size=11)),
                plot_bgcolor='white',
                paper_bgcolor='white',
                margin=dict(l=50, r=50, t=50, b=50)
            )
        )
        pio.templates.default = "research"
    
    def create_publication_trend_chart(self, trends_df, save_path=None):
        """Create publication-quality trend analysis chart"""
        fig = make_subplots(specs=[[{"secondary_y": True}]])
        
        # Main trends
        fig.add_trace(
            go.Scatter(
                x=trends_df['Application_Year'], 
                y=trends_df['Applications'],
                name="Patent Applications", 
                line=dict(color=self.config.COLOR_SCHEME['primary'], width=3),
                marker=dict(size=6, symbol='circle')
            ),
            secondary_y=False,
        )
        
        fig.add_trace(
            go.Scatter(
                x=trends_df['Application_Year'], 
                y=trends_df['Grants'],
                name="Patent Grants", 
                line=dict(color=self.config.COLOR_SCHEME['secondary'], width=3),
                marker=dict(size=6, symbol='square')
            ),
            secondary_y=False,
        )
        
        # Grant rate
        fig.add_trace(
            go.Scatter(
                x=trends_df['Application_Year'], 
                y=trends_df['Grant_Rate'],
                name="Grant Rate (%)", 
                line=dict(color=self.config.COLOR_SCHEME['accent'], width=3, dash='dash'),
                marker=dict(size=6, symbol='diamond')
            ),
            secondary_y=True,
        )
        
        fig.update_layout(
            title=dict(
                text="Temporal Analysis of Patent Applications and Grants",
                x=0.5,
                xanchor='center'
            ),
            xaxis_title="Application Year",
            hovermode="x unified",
            height=500,
            width=800,
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
        )
        
        fig.update_yaxes(title_text="Number of Patents", secondary_y=False)
        fig.update_yaxes(title_text="Grant Rate (%)", secondary_y=True, range=[0, 100])
        
        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))
        
        return fig
    
    def create_technology_landscape_map(self, app_data, grant_data, save_path=None):
        """Create comprehensive technology landscape visualization"""
        # Prepare data
        app_tech = app_data['Cleaned_ICR_Section'].value_counts()
        grant_tech = grant_data['Cleaned_ICR_Section'].value_counts()
        
        tech_df = pd.DataFrame({
            'Technology': app_tech.index,
            'Applications': app_tech.values,
            'Grants': grant_tech.reindex(app_tech.index).fillna(0).values
        })
        tech_df['Grant_Rate'] = (tech_df['Grants'] / tech_df['Applications'] * 100).round(2)
        tech_df['Section_Name'] = tech_df['Technology'].map(self.get_ipc_section_names())
        
        # Create bubble chart
        fig = px.scatter(
            tech_df, 
            x='Applications', 
            y='Grant_Rate',
            size='Applications',
            color='Technology',
            hover_name='Section_Name',
            hover_data={'Applications': True, 'Grants': True, 'Grant_Rate': ':.2f'},
            title="Technology Landscape: Application Volume vs. Grant Success Rate",
            size_max=60,
            color_discrete_sequence=px.colors.qualitative.Set3
        )
        
        fig.update_layout(
            xaxis_title="Number of Applications (Log Scale)",
            yaxis_title="Grant Success Rate (%)",
            xaxis_type="log",
            height=600,
            width=900
        )
        
        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))
        
        return fig
    
    def create_ipc_sunburst_hierarchy(self, df, level='section', save_path=None):
        """Create hierarchical sunburst chart of IPC classification"""
        if level == 'section':
            data = df['Cleaned_ICR_Section'].value_counts()
            path = ['Cleaned_ICR_Section']
        elif level == 'class':
            data = df.groupby(['Cleaned_ICR_Section', 'Cleaned_ICR_Class']).size().reset_index(name='count')
            path = ['Cleaned_ICR_Section', 'Cleaned_ICR_Class']
        else:
            raise ValueError("Level must be 'section' or 'class'")
        
        if level == 'section':
            fig = px.sunburst(
                names=data.index,
                parents=[''] * len(data),
                values=data.values,
                title=f"IPC {level.title()} Distribution",
                color=data.values,
                color_continuous_scale='Viridis'
            )
        else:
            fig = px.sunburst(
                data,
                path=path,
                values='count',
                title=f"IPC {level.title()} Hierarchy",
                color='count',
                color_continuous_scale='Plasma'
            )
        
        fig.update_layout(height=700, width=800)
        
        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))
        
        return fig

    def create_world_choropleth(self, app_df, grant_df, save_path=None):
        """Create world choropleth map for global patent distribution by country"""
        if 'Assignee Country' not in app_df.columns:
            print("Warning: 'Assignee Country' not found in data. Skipping world choropleth.")
            return None

        # Aggregate applications by country
        app_country = app_df['Assignee Country'].value_counts().reset_index()
        app_country.columns = ['country', 'applications']

        # Aggregate grants by country
        grant_country = grant_df['Assignee Country'].value_counts().reset_index()
        grant_country.columns = ['country', 'grants']

        # Merge and calculate grant rate
        country_data = pd.merge(app_country, grant_country, on='country', how='left').fillna(0)
        country_data['grant_rate'] = (country_data['grants'] / country_data['applications'] * 100).round(2)

        # Filter to top countries for better visualization (optional, but helps with clarity)
        top_countries = country_data.nlargest(20, 'applications')

        fig = px.choropleth(
            top_countries,
            locations='country',
            locationmode='country names',
            color='grant_rate',
            hover_name='country',
            hover_data={'applications': True, 'grants': True, 'grant_rate': ':.2f'},
            color_continuous_scale='Viridis',
            title='Global Patent Grant Rates by Assignee Country',
            labels={'grant_rate': 'Grant Rate (%)'}
        )

        fig.update_layout(height=600, width=900)

        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))

        return fig

    def create_geographical_choropleth(self, df, location_col='Assignee State', value_col='Cleaned_ICR_Section', agg_func='count', save_path=None):
        """Create choropleth map for geographical distribution by state (USA fallback)"""
        if location_col not in df.columns:
            print(f"Warning: {location_col} not found in data. Skipping choropleth.")
            return None

        # Aggregate data by location
        geo_data = df.groupby(location_col).size().reset_index(name='count')
        geo_data.columns = ['state', 'count']

        # Ensure state names are full for USA mapping
        state_mapping = {
            'CA': 'California', 'NY': 'New York', 'TX': 'Texas', 'FL': 'Florida',
            'IL': 'Illinois', 'PA': 'Pennsylvania', 'OH': 'Ohio', 'GA': 'Georgia',
            'NC': 'North Carolina', 'MI': 'Michigan'  # Add more as needed
        }
        geo_data['state'] = geo_data['state'].map(lambda x: state_mapping.get(x, x))

        fig = px.choropleth(
            geo_data,
            locations='state',
            locationmode='USA-states',
            color='count',
            color_continuous_scale='Viridis',
            scope='usa',
            title='Geographical Distribution of Patent Applications by State',
            hover_data={'count': True}
        )

        fig.update_layout(height=600, width=900)

        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))

        return fig

    def create_india_choropleth(self, app_df, grant_df, location_col='Assignee State', save_path=None):
        """Create choropleth map for India state-wise patent distribution"""
        if location_col not in app_df.columns:
            print(f"Warning: {location_col} not found in data. Skipping India choropleth.")
            return None

        # Filter for India data (assuming Assignee Country contains 'IN')
        india_app_df = app_df[app_df['Assignee Country'].str.upper().str.contains('IN')]
        india_grant_df = grant_df[grant_df['Assignee Country'].str.upper().str.contains('IN')]

        if india_app_df.empty:
            print("Warning: No India data found. Skipping India choropleth.")
            return None

        # Aggregate applications by state
        india_app_state = india_app_df.groupby(location_col).size().reset_index(name='applications')
        india_app_state.columns = ['state', 'applications']

        # Aggregate grants by state
        india_grant_state = india_grant_df.groupby(location_col).size().reset_index(name='grants')
        india_grant_state.columns = ['state', 'grants']

        # Merge
        india_data = pd.merge(india_app_state, india_grant_state, on='state', how='left').fillna(0)
        india_data['grant_rate'] = (india_data['grants'] / india_data['applications'] * 100).round(2)

        # Map state codes to full names if needed (add more mappings as required)
        state_mapping = {
            'MH': 'Maharashtra', 'KA': 'Karnataka', 'TN': 'Tamil Nadu', 'DL': 'Delhi',
            'GJ': 'Gujarat', 'WB': 'West Bengal', 'UP': 'Uttar Pradesh', 'HR': 'Haryana',
            'AP': 'Andhra Pradesh', 'TS': 'Telangana', 'RJ': 'Rajasthan', 'PN': 'Punjab',
            'KL': 'Kerala', 'OR': 'Odisha', 'MP': 'Madhya Pradesh', 'CG': 'Chhattisgarh',
            'JK': 'Jammu and Kashmir', 'HP': 'Himachal Pradesh', 'UK': 'Uttarakhand',
            'JH': 'Jharkhand', 'BR': 'Bihar', 'AS': 'Assam', 'ML': 'Meghalaya',
            'TR': 'Tripura', 'MZ': 'Mizoram', 'NL': 'Nagaland', 'MN': 'Manipur',
            'AR': 'Arunachal Pradesh', 'GA': 'Goa', 'PY': 'Puducherry', 'CH': 'Chandigarh',
            'DD': 'Daman and Diu', 'DN': 'Dadra and Nagar Haveli', 'LD': 'Lakshadweep'
        }
        india_data['state'] = india_data['state'].map(lambda x: state_mapping.get(x, x))

        try:
            geojson_url = "https://gist.githubusercontent.com/jbrobst/56c13bbbf9d97d187fea01ca62ea67d6/raw/e388b8f20324949f3a5687299fcf03bb4c409551/india_states.geojson"
            response = requests.get(geojson_url)
            response.raise_for_status()
            geojson = response.json()

            fig = px.choropleth(
                india_data,
                geojson=geojson,
                locations='state',
                featureidkey="properties.ST_NM",
                color='grant_rate',
                hover_name='state',
                hover_data={'applications': True, 'grants': True, 'grant_rate': ':.2f'},
                color_continuous_scale='Viridis',
                scope='asia',
                title='India State-wise Patent Grant Rates',
                labels={'grant_rate': 'Grant Rate (%)'}
            )
        except Exception as e:
            print(f"Error creating India choropleth: {e}. Skipping.")
            return None

        fig.update_layout(height=600, width=900)

        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))

        return fig

    def create_india_technology_choropleth(self, app_df, grant_df, location_col='Assignee State', save_path=None):
        """Create choropleth map for India state-wise dominant technology (IPC Section)"""
        if location_col not in app_df.columns or 'Cleaned_ICR_Section' not in app_df.columns:
            print(f"Warning: {location_col} or Cleaned_ICR_Section not found in data. Skipping India technology choropleth.")
            return None

        # Filter for India data
        india_app_df = app_df[app_df['Assignee Country'].str.upper().str.contains('IN')]

        if india_app_df.empty:
            print("Warning: No India data found. Skipping India technology choropleth.")
            return None

        # Find dominant IPC section per state
        state_tech = india_app_df.groupby([location_col, 'Cleaned_ICR_Section']).size().reset_index(name='count')
        dominant_tech = state_tech.loc[state_tech.groupby(location_col)['count'].idxmax()].reset_index(drop=True)
        dominant_tech.columns = ['state', 'dominant_tech', 'count']

        # Map state codes to full names
        state_mapping = {
            'MH': 'Maharashtra', 'KA': 'Karnataka', 'TN': 'Tamil Nadu', 'DL': 'Delhi',
            'GJ': 'Gujarat', 'WB': 'West Bengal', 'UP': 'Uttar Pradesh', 'HR': 'Haryana',
            'AP': 'Andhra Pradesh', 'TS': 'Telangana', 'RJ': 'Rajasthan', 'PN': 'Punjab',
            'KL': 'Kerala', 'OR': 'Odisha', 'MP': 'Madhya Pradesh', 'CG': 'Chhattisgarh',
            'JK': 'Jammu and Kashmir', 'HP': 'Himachal Pradesh', 'UK': 'Uttarakhand',
            'JH': 'Jharkhand', 'BR': 'Bihar', 'AS': 'Assam', 'ML': 'Meghalaya',
            'TR': 'Tripura', 'MZ': 'Mizoram', 'NL': 'Nagaland', 'MN': 'Manipur',
            'AR': 'Arunachal Pradesh', 'GA': 'Goa', 'PY': 'Puducherry', 'CH': 'Chandigarh',
            'DD': 'Daman and Diu', 'DN': 'Dadra and Nagar Haveli', 'LD': 'Lakshadweep'
        }
        dominant_tech['state'] = dominant_tech['state'].map(lambda x: state_mapping.get(x, x))

        # Add section names
        dominant_tech['tech_name'] = dominant_tech['dominant_tech'].map(self.ipc_sections)

        try:
            geojson_url = "https://gist.githubusercontent.com/jbrobst/56c13bbbf9d97d187fea01ca62ea67d6/raw/e388b8f20324949f3a5687299fcf03bb4c409551/india_states.geojson"
            response = requests.get(geojson_url)
            response.raise_for_status()
            geojson = response.json()

            fig = px.choropleth(
                dominant_tech,
                geojson=geojson,
                locations='state',
                featureidkey="properties.ST_NM",
                color='dominant_tech',
                hover_name='state',
                hover_data={'dominant_tech': True, 'tech_name': True, 'count': True},
                color_discrete_sequence=px.colors.qualitative.Set3,
                scope='asia',
                title='India State-wise Dominant Technology (IPC Section)',
                labels={'dominant_tech': 'IPC Section'}
            )
        except Exception as e:
            print(f"Error creating India technology choropleth: {e}. Skipping.")
            return None

        fig.update_layout(height=600, width=900)

        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))

        return fig

    def create_india_industry_choropleth(self, app_df, grant_df, location_col='Assignee State', save_path=None):
        """Create choropleth map for India state-wise dominant industry"""
        if location_col not in app_df.columns or 'Industry' not in app_df.columns:
            print(f"Warning: {location_col} or Industry not found in data. Skipping India industry choropleth.")
            return None

        # Filter for India data
        india_app_df = app_df[app_df['Assignee Country'].str.upper().str.contains('IN')]

        if india_app_df.empty:
            print("Warning: No India data found. Skipping India industry choropleth.")
            return None

        # Find dominant industry per state
        state_industry = india_app_df.groupby([location_col, 'Industry']).size().reset_index(name='count')
        dominant_industry = state_industry.loc[state_industry.groupby(location_col)['count'].idxmax()].reset_index(drop=True)
        dominant_industry.columns = ['state', 'dominant_industry', 'count']

        # Map state codes to full names
        state_mapping = {
            'MH': 'Maharashtra', 'KA': 'Karnataka', 'TN': 'Tamil Nadu', 'DL': 'Delhi',
            'GJ': 'Gujarat', 'WB': 'West Bengal', 'UP': 'Uttar Pradesh', 'HR': 'Haryana',
            'AP': 'Andhra Pradesh', 'TS': 'Telangana', 'RJ': 'Rajasthan', 'PN': 'Punjab',
            'KL': 'Kerala', 'OR': 'Odisha', 'MP': 'Madhya Pradesh', 'CG': 'Chhattisgarh',
            'JK': 'Jammu and Kashmir', 'HP': 'Himachal Pradesh', 'UK': 'Uttarakhand',
            'JH': 'Jharkhand', 'BR': 'Bihar', 'AS': 'Assam', 'ML': 'Meghalaya',
            'TR': 'Tripura', 'MZ': 'Mizoram', 'NL': 'Nagaland', 'MN': 'Manipur',
            'AR': 'Arunachal Pradesh', 'GA': 'Goa', 'PY': 'Puducherry', 'CH': 'Chandigarh',
            'DD': 'Daman and Diu', 'DN': 'Dadra and Nagar Haveli', 'LD': 'Lakshadweep'
        }
        dominant_industry['state'] = dominant_industry['state'].map(lambda x: state_mapping.get(x, x))

        try:
            geojson_url = "https://gist.githubusercontent.com/jbrobst/56c13bbbf9d97d187fea01ca62ea67d6/raw/e388b8f20324949f3a5687299fcf03bb4c409551/india_states.geojson"
            response = requests.get(geojson_url)
            response.raise_for_status()
            geojson = response.json()

            fig = px.choropleth(
                dominant_industry,
                geojson=geojson,
                locations='state',
                featureidkey="properties.ST_NM",
                color='dominant_industry',
                hover_name='state',
                hover_data={'dominant_industry': True, 'count': True},
                color_discrete_sequence=px.colors.qualitative.Set2,
                scope='asia',
                title='India State-wise Dominant Industry',
                labels={'dominant_industry': 'Industry'}
            )
        except Exception as e:
            print(f"Error creating India industry choropleth: {e}. Skipping.")
            return None

        fig.update_layout(height=600, width=900)

        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))

        return fig

    def create_sankey_flow(self, app_df, grant_df, save_path=None):
        """Create Sankey diagram for flows between IPC sections and grant status"""
        # Prepare nodes: IPC sections for apps and grants
        app_sections = app_df['Cleaned_ICR_Section'].value_counts().to_dict()
        grant_sections = grant_df['Cleaned_ICR_Section'].value_counts().to_dict()
        
        # Unique sections
        all_sections = sorted(set(list(app_sections.keys()) + list(grant_sections.keys())))
        
        # Nodes: left for applications, right for grants
        labels = [f"App_{sec}" for sec in all_sections] + [f"Grant_{sec}" for sec in all_sections]
        
        # Links: from app sections to grant sections (simplified: assume flow based on counts)
        source = []
        target = []
        value = []
        
        for sec in all_sections:
            app_count = app_sections.get(sec, 0)
            grant_count = grant_sections.get(sec, 0)
            if app_count > 0:
                source.append(all_sections.index(sec))
                target.append(len(all_sections) + all_sections.index(sec))
                value.append(min(app_count, grant_count))  # Simplified flow
        
        fig = go.Figure(data=[go.Sankey(
            node=dict(
                pad=15,
                thickness=20,
                line=dict(color="black", width=0.5),
                label=labels,
                color=[self.config.COLOR_SCHEME['primary']] * len(labels)
            ),
            link=dict(
                source=source,
                target=target,
                value=value
            )
        )])
        
        fig.update_layout(title_text="Sankey Diagram: Flow from Applications to Grants by IPC Section", height=700, width=900)
        
        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))
        
        return fig

    def create_treemap_hierarchy(self, df, save_path=None):
        """Create treemap for IPC hierarchical data"""
        # Prepare hierarchical data, filter out None values
        hierarchy_df = df.groupby(['Cleaned_ICR_Section', 'Cleaned_ICR_Class']).size().reset_index(name='count')
        hierarchy_df = hierarchy_df.dropna(subset=['Cleaned_ICR_Section', 'Cleaned_ICR_Class'])
        hierarchy_df['Cleaned_ICR_Section_Name'] = hierarchy_df['Cleaned_ICR_Section'].map(self.ipc_sections)

        if hierarchy_df.empty:
            print("Warning: No valid hierarchical data for treemap. Skipping.")
            return None

        fig = px.treemap(
            hierarchy_df,
            path=['Cleaned_ICR_Section_Name', 'Cleaned_ICR_Section', 'Cleaned_ICR_Class'],
            values='count',
            color='count',
            color_continuous_scale='Plasma',
            title='IPC Hierarchy Treemap'
        )

        fig.update_layout(height=700, width=900)

        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))

        return fig

    def create_tech_hierarchy_sunburst(self, app_df, grant_df, save_path=None):
        """Create hierarchical sunburst for Tech Domain > Tech Sub Domain > IPC, colored by grant rate"""
        required_cols = ['Tech Domain', 'Tech Sub Domain', 'Cleaned_ICR_Section', 'Application No.']
        if not all(col in app_df.columns for col in required_cols) or not all(col in grant_df.columns for col in ['Application No.', 'Cleaned_ICR_Section']):
            print("Warning: Required columns missing for tech hierarchy sunburst. Skipping.")
            return None

        # Merge to get grant status
        merged = pd.merge(
            app_df[required_cols],
            grant_df[['Application No.', 'Cleaned_ICR_Section']],
            on='Application No.',
            how='left',
            suffixes=('_app', '_grant')
        )
        merged['Granted'] = ~merged['Cleaned_ICR_Section_grant'].isna()

        # Group by hierarchy and aggregate
        hierarchy_agg = merged.groupby(['Tech Domain', 'Tech Sub Domain', 'Cleaned_ICR_Section_app']).agg({
            'Granted': ['count', 'sum']
        }).round(2)
        hierarchy_agg.columns = ['count', 'grants']
        hierarchy_agg['grant_rate'] = (hierarchy_agg['grants'] / hierarchy_agg['count'] * 100).round(2)
        hierarchy_agg = hierarchy_agg.reset_index()
        hierarchy_agg = hierarchy_agg.rename(columns={'Cleaned_ICR_Section_app': 'Cleaned_ICR_Section'})

        # Filter out any NaN/None in hierarchy
        hierarchy_agg = hierarchy_agg.dropna(subset=['Tech Domain', 'Tech Sub Domain', 'Cleaned_ICR_Section'])

        if hierarchy_agg.empty:
            print("Warning: No valid hierarchical data for tech sunburst. Skipping.")
            return None

        # Create sunburst
        fig = px.sunburst(
            hierarchy_agg,
            path=['Tech Domain', 'Tech Sub Domain', 'Cleaned_ICR_Section'],
            values='count',
            color='grant_rate',
            color_continuous_scale='RdYlGn',  # Red-Yellow-Green for grant rate (low to high)
            title='Technology Hierarchy Sunburst: Colored by Grant Rate',
            hover_data={'count': True, 'grants': True, 'grant_rate': ':.2f'}
        )

        fig.update_layout(height=700, width=800)

        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))

        return fig

    def create_radar_comparison(self, app_df, grant_df, top_n=5, save_path=None):
        """Create radar chart for multi-dimensional comparison of top tech domains"""
        # Use Tech Domain if available, else IPC Section
        domain_col = 'Tech Domain' if 'Tech Domain' in app_df.columns else 'Cleaned_ICR_Section'

        top_domains = app_df[domain_col].value_counts().head(top_n).index
        
        # Prepare data for each domain
        domains_data = []
        for domain in top_domains:
            app_count = len(app_df[app_df[domain_col] == domain])
            grant_count = len(grant_df[grant_df[domain_col] == domain])
            grant_rate = (grant_count / app_count * 100) if app_count > 0 else 0
            hhi_contrib = (app_count / len(app_df)) ** 2 * 10000  # Simplified
            
            domains_data.append({
                'domain': domain,
                'applications': app_count,
                'grants': grant_count,
                'grant_rate': grant_rate,
                'concentration': hhi_contrib
            })
        
        domains_df = pd.DataFrame(domains_data)
        
        # Normalize for radar
        for col in ['applications', 'grants', 'grant_rate', 'concentration']:
            domains_df[col + '_norm'] = (domains_df[col] - domains_df[col].min()) / (domains_df[col].max() - domains_df[col].min())
        
        # Create radar for each domain
        fig = go.Figure()
        
        angles = ['applications_norm', 'grants_norm', 'grant_rate_norm', 'concentration_norm']
        angle_labels = ['Applications', 'Grants', 'Grant Rate', 'Concentration']
        
        for idx, domain in enumerate(top_domains):
            values = domains_df.loc[idx, angles].tolist()
            values += values[:1]  # Complete the circle
            fig.add_trace(go.Scatterpolar(
                r=values,
                theta=angle_labels + [angle_labels[0]],
                fill='toself',
                name=domain
            ))
        
        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, 1]
                )),
            showlegend=True,
            title="Radar Chart: Multi-dimensional Comparison of Top Technology Domains",
            height=600,
            width=800
        )
        
        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))
        
        return fig

    def create_network_graph(self, df, node_col='Tech Domain', edge_col='Tech Sub Domain', save_path=None):
        """Create network graph for technology clustering analysis"""
        if node_col not in df.columns or edge_col not in df.columns:
            print(f"Warning: {node_col} or {edge_col} not found. Skipping network graph.")
            return None
        
        G = nx.Graph()
        
        # Add nodes (domains) with degree/ count as attribute
        domains = df[node_col].value_counts()
        for domain in domains.index:
            G.add_node(domain, size=domains[domain])
        
        # Calculate weighted edges based on co-occurrence counts
        edge_counts = df.groupby([node_col, edge_col]).size().reset_index(name='weight')
        for _, row in edge_counts.iterrows():
            G.add_edge(row[node_col], row[edge_col], weight=row['weight'])
        
        # Position nodes
        pos = nx.spring_layout(G, k=3, iterations=50)  # Adjusted for better layout
        
        # Create Plotly figure with weighted edges
        edge_x = []
        edge_y = []
        edge_widths = []
        for edge in G.edges(data=True):
            x0, y0 = pos[edge[0]]
            x1, y1 = pos[edge[1]]
            edge_x.extend([x0, x1, None])
            edge_y.extend([y0, y1, None])
            # Normalize weights for line width (0.5 to 5)
            weight = edge[2]['weight']
            norm_width = 0.5 + (weight / edge_counts['weight'].max()) * 4.5
            edge_widths.extend([norm_width, norm_width, None])
        
        # Since Plotly Scatter can't have varying widths per segment easily, use average or multiple traces if needed
        # For simplicity, use a single trace with fixed width; advanced would require multiple traces
        avg_width = np.mean([d['weight'] for _, _, d in G.edges(data=True)]) / edge_counts['weight'].max() * 5
        edge_trace = go.Scatter(
            x=edge_x, 
            y=edge_y, 
            line=dict(width=avg_width, color='#888'), 
            hoverinfo='none', 
            mode='lines'
        )
        
        # Nodes with size based on degree/count
        node_x = [pos[node][0] for node in G.nodes()]
        node_y = [pos[node][1] for node in G.nodes()]
        node_sizes = [G.nodes[node].get('size', 10) / df[node_col].nunique() * 20 for node in G.nodes()]  # Normalize size
        node_trace = go.Scatter(
            x=node_x, 
            y=node_y,
            mode='markers+text',
            hoverinfo='text',
            text=list(G.nodes()),
            textposition="middle center",
            marker=dict(
                size=node_sizes, 
                color='LightSkyBlue',
                line=dict(width=2, color='DarkSlateGrey')
            )
        )
        
        fig = go.Figure(data=[edge_trace, node_trace],
                       layout=go.Layout(
                           title='Network Graph: Technology Clustering (Weighted Edges)',
                           showlegend=False,
                           hovermode='closest',
                           margin=dict(b=20, l=5, r=5, t=40),
                           annotations=[dict(
                               activeshape=False, 
                               xref='paper', 
                               yref='paper', 
                               x=0.5, 
                               y=-0.1,
                               xanchor='center', 
                               yanchor='top',
                               text='Technology Domain (Nodes) - Sub Domain (Edges) Network'
                           )],
                           xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                           yaxis=dict(showgrid=False, zeroline=False, showticklabels=False)
                       ))
        
        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))
        
        return fig
    
    def create_industry_treemap(self, app_df, grant_df, save_path=None):
        """Create treemap for industry distribution, sized by applications, colored by grant rate"""
        if 'Industry' not in app_df.columns:
            print("Warning: 'Industry' not found in data. Skipping industry treemap.")
            return None

        # Aggregate applications by industry
        industry_apps = app_df['Industry'].value_counts().reset_index()
        industry_apps.columns = ['Industry', 'Applications']

        # Aggregate grants by industry
        if 'Industry' in grant_df.columns:
            industry_grants = grant_df['Industry'].value_counts().reset_index()
            industry_grants.columns = ['Industry', 'Grants']
            industry_data = pd.merge(industry_apps, industry_grants, on='Industry', how='left').fillna(0)
        else:
            industry_data = industry_apps
            industry_data['Grants'] = 0

        industry_data['Grant_Rate'] = (industry_data['Grants'] / industry_data['Applications'] * 100).round(2)

        # Filter to top industries for better visualization
        top_industries = industry_data.nlargest(20, 'Applications')

        fig = px.treemap(
            top_industries,
            path=['Industry'],
            values='Applications',
            color='Grant_Rate',
            color_continuous_scale='RdYlGn',
            title='Industry Distribution Treemap: Sized by Applications, Colored by Grant Rate',
            hover_data={'Applications': True, 'Grants': True, 'Grant_Rate': ':.2f'}
        )

        fig.update_layout(height=700, width=900)

        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))

        return fig

    def create_industry_sankey(self, app_df, grant_df, save_path=None):
        """Create Sankey diagram for flow from Industry to Tech Domain or grant status"""
        required_cols = ['Industry', 'Tech Domain', 'Application No.']
        if not all(col in app_df.columns for col in required_cols):
            print("Warning: Required columns missing for industry sankey. Skipping.")
            return None

        # Merge to get grant status
        merged = pd.merge(
            app_df[required_cols],
            grant_df[['Application No.']],
            on='Application No.',
            how='left',
            indicator=True
        )
        merged['Granted'] = (merged['_merge'] == 'both').astype(int)

        # Aggregate flows: Industry -> Tech Domain -> Grant Status
        flow_data = merged.groupby(['Industry', 'Tech Domain', 'Granted']).size().reset_index(name='count')

        # Filter to top industries and tech domains for clarity
        top_industries = merged['Industry'].value_counts().head(10).index
        top_tech_domains = merged['Tech Domain'].value_counts().head(10).index
        flow_data = flow_data[
            (flow_data['Industry'].isin(top_industries)) &
            (flow_data['Tech Domain'].isin(top_tech_domains))
        ]

        if flow_data.empty:
            print("Warning: No valid flow data for industry sankey. Skipping.")
            return None

        # Create nodes: Industries + Tech Domains + Grant Status
        industries = flow_data['Industry'].unique()
        tech_domains = flow_data['Tech Domain'].unique()
        grant_statuses = ['Not Granted', 'Granted']

        labels = list(industries) + list(tech_domains) + grant_statuses

        # Create links
        source = []
        target = []
        value = []

        for _, row in flow_data.iterrows():
            # Industry to Tech Domain
            source.append(list(industries).index(row['Industry']))
            target.append(len(industries) + list(tech_domains).index(row['Tech Domain']))
            value.append(row['count'])

            # Tech Domain to Grant Status
            target.append(len(industries) + len(tech_domains) + int(row['Granted']))
            source.append(len(industries) + list(tech_domains).index(row['Tech Domain']))
            value.append(row['count'])

        fig = go.Figure(data=[go.Sankey(
            node=dict(
                pad=15,
                thickness=20,
                line=dict(color="black", width=0.5),
                label=labels,
                color=[self.config.COLOR_SCHEME['primary']] * len(labels)
            ),
            link=dict(
                source=source,
                target=target,
                value=value
            )
        )])

        fig.update_layout(
            title_text="Sankey Diagram: Flow from Industry to Tech Domain to Grant Status",
            height=700,
            width=900
        )

        if save_path:
            fig.write_image(save_path, scale=2, format='png')
            fig.write_html(save_path.replace('.png', '.html'))

        return fig

    def get_ipc_section_names(self):
        """Return IPC section names for labeling"""
        return self.ipc_sections

# Research report generator
class ResearchReportGenerator:
    def __init__(self, config, analyzer, visualizer):
        self.config = config
        self.analyzer = analyzer
        self.visualizer = visualizer
    
    def generate_comprehensive_report(self, app_df, grant_df, metrics):
        """Generate comprehensive research report"""
        report = {
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'metrics': metrics,
            'figures': {},
            'tables': {},
            'insights': []
        }
        
        # Generate key visualizations
        trends = self.analyzer.get_trend_analysis()
        report['figures']['trend_analysis'] = self.visualizer.create_publication_trend_chart(
            trends, 
            f"{self.config.FIGURES_DIR}/trend_analysis.png"
        )
        
        if 'Cleaned_ICR_Section' in app_df.columns and 'Cleaned_ICR_Section' in grant_df.columns:
            report['figures']['technology_landscape'] = self.visualizer.create_technology_landscape_map(
                app_df, grant_df, 
                f"{self.config.FIGURES_DIR}/technology_landscape.png"
            )
            
            report['figures']['ipc_sunburst'] = self.visualizer.create_ipc_sunburst_hierarchy(
                app_df, 
                'section', 
                f"{self.config.FIGURES_DIR}/ipc_sunburst.png"
            )
            
            report['figures']['ipc_treemap'] = self.visualizer.create_treemap_hierarchy(
                app_df,
                f"{self.config.FIGURES_DIR}/ipc_treemap.png"
            )
            
            report['figures']['sankey_flow'] = self.visualizer.create_sankey_flow(
                app_df, grant_df,
                f"{self.config.FIGURES_DIR}/sankey_flow.png"
            )
            
            report['figures']['radar_comparison'] = self.visualizer.create_radar_comparison(
                app_df, grant_df, top_n=5,
                save_path=f"{self.config.FIGURES_DIR}/radar_comparison.png"
            )
            
            report['figures']['network_graph'] = self.visualizer.create_network_graph(
                app_df,
                f"{self.config.FIGURES_DIR}/network_graph.png"
            )

            report['figures']['tech_hierarchy_sunburst'] = self.visualizer.create_tech_hierarchy_sunburst(
                app_df, grant_df,
                f"{self.config.FIGURES_DIR}/tech_hierarchy_sunburst.png"
            )
        
        # Geographical choropleth: world map if country data available, else USA state fallback
        if 'Assignee Country' in app_df.columns:
            report['figures']['world_choropleth'] = self.visualizer.create_world_choropleth(
                app_df, grant_df,
                f"{self.config.FIGURES_DIR}/world_choropleth.png"
            )
            # Add India state-wise choropleth if Assignee State data available
            report['figures']['india_choropleth'] = self.visualizer.create_india_choropleth(
                app_df, grant_df,
                save_path=f"{self.config.FIGURES_DIR}/india_choropleth.png"
            )
            report['figures']['india_technology_choropleth'] = self.visualizer.create_india_technology_choropleth(
                app_df, grant_df,
                save_path=f"{self.config.FIGURES_DIR}/india_technology_choropleth.png"
            )
            report['figures']['india_industry_choropleth'] = self.visualizer.create_india_industry_choropleth(
                app_df, grant_df,
                save_path=f"{self.config.FIGURES_DIR}/india_industry_choropleth.png"
            )
        elif 'Assignee State' in app_df.columns:
            report['figures']['geographical_choropleth'] = self.visualizer.create_geographical_choropleth(
                app_df,
                save_path=f"{self.config.FIGURES_DIR}/geographical_choropleth.png"
            )

        # Industry-wise visualizations
        if 'Industry' in app_df.columns:
            report['figures']['industry_treemap'] = self.visualizer.create_industry_treemap(
                app_df, grant_df,
                f"{self.config.FIGURES_DIR}/industry_treemap.png"
            )

            if 'Tech Domain' in app_df.columns:
                report['figures']['industry_sankey'] = self.visualizer.create_industry_sankey(
                    app_df, grant_df,
                    f"{self.config.FIGURES_DIR}/industry_sankey.png"
                )
        
        # Generate insights
        report['insights'] = self.generate_research_insights(metrics, trends, app_df, grant_df)
        
        # Save data tables
        self.save_research_tables(app_df, grant_df, trends)
        
        # Generate HTML report
        self.generate_html_report(report)
        
        # Generate DOCX report
        self.generate_docx_report(report)
        
        return report
    
    def generate_research_insights(self, metrics, trends, app_df, grant_df):
        """Generate research insights from analysis"""
        insights = []

        # Insight 1: Overall grant rate
        grant_rate = metrics['overall_grant_rate']
        insights.append({
            'title': 'Overall Patent Grant Efficiency',
            'content': f'The overall grant rate for patents is {grant_rate:.1f}%, indicating {("moderate" if grant_rate > 50 else "low")} conversion efficiency from application to grant.'
        })

        # Insight 2: Technology concentration
        hhi = metrics['technology_concentration']['hhi']
        concentration_level = "highly concentrated" if hhi > 2500 else "moderately concentrated" if hhi > 1500 else "diversified"
        insights.append({
            'title': 'Technology Concentration Analysis',
            'content': f'The technology landscape shows {concentration_level} distribution (HHI: {hhi:.0f}), with certain IPC sections dominating the patent landscape.'
        })

        # Insight 3: Temporal trends
        recent_trend = trends.iloc[-3:]  # Last 3 years
        avg_growth = recent_trend['Applications'].pct_change().mean() * 100
        insights.append({
            'title': 'Recent Growth Patterns',
            'content': f'Recent years show an average annual growth rate of {avg_growth:.1f}% in patent applications, indicating {"increasing" if avg_growth > 0 else "decreasing"} innovation activity.'
        })

        # Insight 4: Industry concentration (if Industry data available)
        if 'Industry' in app_df.columns:
            industry_counts = app_df['Industry'].value_counts()
            industry_hhi = (industry_counts / industry_counts.sum()) ** 2 * 10000
            industry_hhi_total = industry_hhi.sum()
            industry_concentration_level = "highly concentrated" if industry_hhi_total > 2500 else "moderately concentrated" if industry_hhi_total > 1500 else "diversified"
            top_industry = industry_counts.index[0] if len(industry_counts) > 0 else "Unknown"
            insights.append({
                'title': 'Industry Concentration Analysis',
                'content': f'The industry distribution shows {industry_concentration_level} concentration (HHI: {industry_hhi_total:.0f}), with {top_industry} leading in patent activity.'
            })

        # Insight 5: Top countries by grants (if country data available)
        if 'Assignee Country' in grant_df.columns:
            top_countries = grant_df['Assignee Country'].value_counts().head(3)
            top_country = top_countries.index[0] if len(top_countries) > 0 else "Unknown"
            top_country_grants = top_countries.iloc[0] if len(top_countries) > 0 else 0
            insights.append({
                'title': 'Geographical Leadership in Patent Grants',
                'content': f'{top_country} leads in patent grants with {top_country_grants:,} granted patents, representing the most active country in successful patent outcomes.'
            })

        return insights
    
    def save_research_tables(self, app_df, grant_df, trends):
        """Save research data tables"""
        # Top technology areas
        if 'Cleaned_ICR_Section' in app_df.columns:
            top_tech_apps = app_df['Cleaned_ICR_Section'].value_counts().head(10)
            if 'Cleaned_ICR_Section' in grant_df.columns:
                top_tech_grants = grant_df['Cleaned_ICR_Section'].value_counts().head(10)
                tech_table = pd.DataFrame({
                    'IPC Section': top_tech_apps.index,
                    'Applications': top_tech_apps.values,
                    'Grants': top_tech_grants.reindex(top_tech_apps.index).fillna(0).values
                })
                tech_table['Grant Rate (%)'] = (tech_table['Grants'] / tech_table['Applications'] * 100).round(2)
            else:
                tech_table = pd.DataFrame({
                    'IPC Section': top_tech_apps.index,
                    'Applications': top_tech_apps.values
                })
            tech_table.to_csv(f"{self.config.DATA_DIR}/top_technology_areas.csv", index=False)
        trends.to_csv(f"{self.config.DATA_DIR}/temporal_trends.csv", index=False)
    
    def generate_html_report(self, report):
        """Generate interactive HTML research report"""
        html_content = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Overall Patent Analysis - Research Report</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                .header {{ text-align: center; border-bottom: 2px solid #2E8B57; padding-bottom: 20px; }}
                .section {{ margin: 40px 0; }}
                .insight {{ background: #f8f9fa; padding: 15px; border-left: 4px solid #2E8B57; margin: 10px 0; }}
                .metric-card {{ background: white; padding: 20px; margin: 10px; border-radius: 5px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
                .metrics-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 20px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Overall Patent Analysis</h1>
                <p>Research Report Generated on {report['timestamp']}</p>
            </div>
            
            <div class="section">
                <h2>Executive Summary</h2>
                <div class="metrics-grid">
                    <div class="metric-card">
                        <h3>Total Applications</h3>
                        <p style="font-size: 24px; color: #2E8B57;">{report['metrics']['total_applications']:,}</p>
                    </div>
                    <div class="metric-card">
                        <h3>Total Grants</h3>
                        <p style="font-size: 24px; color: #1F77B4;">{report['metrics']['total_grants']:,}</p>
                    </div>
                    <div class="metric-card">
                        <h3>Grant Rate</h3>
                        <p style="font-size: 24px; color: #D62728;">{report['metrics']['overall_grant_rate']:.1f}%</p>
                    </div>
                </div>
            </div>
            
            <div class="section">
                <h2>Key Insights</h2>
                {"".join([f'<div class="insight"><h3>{insight["title"]}</h3><p>{insight["content"]}</p></div>' for insight in report['insights']])}
            </div>
            
            <div class="section">
                <h2>Visualizations</h2>
                <p><strong>Note:</strong> Interactive visualizations are saved as separate HTML files in the output directory.</p>
                <div>
                    <h3>Temporal Trend Analysis</h3>
                    <img src="figures/trend_analysis.png" alt="Trend Analysis" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>Technology Landscape</h3>
                    <img src="figures/technology_landscape.png" alt="Technology Landscape" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>IPC Sunburst</h3>
                    <img src="figures/ipc_sunburst.png" alt="IPC Sunburst" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>IPC Treemap</h3>
                    <img src="figures/ipc_treemap.png" alt="IPC Treemap" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>Sankey Flow</h3>
                    <img src="figures/sankey_flow.png" alt="Sankey Flow" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>Radar Comparison</h3>
                    <img src="figures/radar_comparison.png" alt="Radar Comparison" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>Network Graph</h3>
                    <img src="figures/network_graph.png" alt="Network Graph" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>World Choropleth</h3>
                    <img src="figures/world_choropleth.png" alt="World Choropleth" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>India Choropleth</h3>
                    <img src="figures/india_choropleth.png" alt="India Choropleth" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>India Technology Choropleth</h3>
                    <img src="figures/india_technology_choropleth.png" alt="India Technology Choropleth" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>India Industry Choropleth</h3>
                    <img src="figures/india_industry_choropleth.png" alt="India Industry Choropleth" style="width: 100%; max-width: 800px;">
                </div>
                <div>
                    <h3>Geographical Choropleth</h3>
                    <img src="figures/geographical_choropleth.png" alt="Geographical Choropleth" style="width: 100%; max-width: 800px;">
                </div>
            </div>
        </body>
        </html>
        """
        
        with open(f"{self.config.HTML_DIR}/research_report.html", 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def generate_docx_report(self, report):
        """Generate DOCX research report with summary and embedded visualizations"""
        try:
            from docx import Document
            from docx.shared import Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import os
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Overall Patent Analysis Research Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Timestamp
            doc.add_paragraph(f"Generated on: {report['timestamp']}")
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            summary_table = doc.add_table(rows=1, cols=3)
            summary_table.style = 'Table Grid'
            
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            hdr_cells[2].text = 'Description'
            
            row_cells = summary_table.add_row().cells
            row_cells[0].text = 'Total Applications'
            row_cells[1].text = str(report['metrics']['total_applications'])
            row_cells[2].text = 'Total patent applications analyzed'
            
            row_cells = summary_table.add_row().cells
            row_cells[0].text = 'Total Grants'
            row_cells[1].text = str(report['metrics']['total_grants'])
            row_cells[2].text = 'Total granted patents'
            
            row_cells = summary_table.add_row().cells
            row_cells[0].text = 'Overall Grant Rate'
            row_cells[1].text = f"{report['metrics']['overall_grant_rate']:.1f}%"
            row_cells[2].text = 'Percentage of applications granted'
            
            # Key Insights
            doc.add_heading('Key Insights', level=1)
            for insight in report['insights']:
                doc.add_heading(insight['title'], level=2)
                doc.add_paragraph(insight['content'])
            
            # Embed images if they exist
            doc.add_heading('Visualizations', level=1)
            for fig_name, fig in report['figures'].items():
                img_path = f"{self.config.FIGURES_DIR}/{fig_name}.png"
                if os.path.exists(img_path):
                    doc.add_heading(fig_name.replace('_', ' ').title(), level=2)
                    doc.add_picture(img_path, width=Inches(6))
                else:
                    doc.add_paragraph(f"Visualization '{fig_name}' not available.")
            
            # Save DOCX
            docx_path = f"{self.config.HTML_DIR}/research_report.docx"
            doc.save(docx_path)
            print(f"DOCX report saved to: {docx_path}")
            
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            print("DOCX report generation skipped.")
        except Exception as e:
            print(f"Error generating DOCX report: {e}")

# Patent analyzer class (extracted from notebook)
class PatentAnalyzer:
    """Advanced patent data analysis"""
    
    def __init__(self, app_df, grant_df):
        self.app_df = app_df
        self.grant_df = grant_df
    
    def get_trend_analysis(self):
        """Analyze trends over time"""
        app_trends = self.app_df.groupby('Application_Year').size().reset_index(name='Applications')
        grant_trends = self.grant_df.groupby('Application_Year').size().reset_index(name='Grants')
        
        trends = pd.merge(app_trends, grant_trends, on='Application_Year', how='outer').fillna(0)
        trends['Grant_Rate'] = (trends['Grants'] / trends['Applications'] * 100).round(2)
        
        return trends
    
    def analyze_ipc_distribution(self, df, ipc_level='Cleaned_ICR_Section'):
        """Analyze IPC code distribution"""
        return df[ipc_level].value_counts().head(15)
    
    def calculate_grant_rates_by_technology(self):
        """Calculate grant rates by technology area"""
        # Merge applications and grants
        merged = pd.merge(
            self.app_df[['Application No.', 'Cleaned_ICR_Section']],
            self.grant_df[['Application No.', 'Cleaned_ICR_Section']],
            on='Application No.',
            how='left',
            suffixes=('_app', '_grant')
        )

        merged['Granted'] = ~merged['Cleaned_ICR_Section_grant'].isna()

        grant_rates = merged.groupby('Cleaned_ICR_Section_app').agg({
            'Granted': ['count', 'sum']
        }).round(2)

        grant_rates.columns = ['Total_Applications', 'Grants']
        grant_rates['Grant_Rate'] = (grant_rates['Grants'] / grant_rates['Total_Applications'] * 100).round(2)

        return grant_rates.sort_values('Total_Applications', ascending=False)

# Main execution block
def main():
    print("=== Overall Patent Analysis ===\n")

    # Load data
    print("1. Loading and preprocessing data...")
    green_app_df = pd.read_csv(config.DATA_PATHS['applications'])
    green_grant_df = pd.read_csv(config.DATA_PATHS['grants'])

    # Create Application_Year from Filing/Application Date
    green_app_df['Application_Year'] = pd.to_datetime(green_app_df['Filing/Application Date'], errors='coerce').dt.year
    green_grant_df['Application_Year'] = pd.to_datetime(green_grant_df['Filing/Application Date'], errors='coerce').dt.year

    # Enhanced preprocessing
    preprocessor = ResearchPatentPreprocessor(green_app_df, green_grant_df)
    green_app_df = preprocessor.enhanced_ipc_analysis(green_app_df)
    green_grant_df = preprocessor.enhanced_ipc_analysis(green_grant_df)

    # Calculate research metrics
    metrics = preprocessor.calculate_research_metrics()

    # Initialize analyzer and visualizer
    analyzer = PatentAnalyzer(green_app_df, green_grant_df)  # Your existing analyzer
    visualizer = ResearchVisualizer(config)
    report_generator = ResearchReportGenerator(config, analyzer, visualizer)

    # Generate comprehensive report
    print("2. Generating research report...")
    report = report_generator.generate_comprehensive_report(green_app_df, green_grant_df, metrics)

    print("3. Analysis completed successfully!")
    print(f"\n=== RESEARCH OUTPUTS ===")
    print(f"✓ Figures saved to: {config.FIGURES_DIR}/")
    print(f"✓ Data tables saved to: {config.DATA_DIR}/")
    print(f"✓ HTML report saved to: {config.HTML_DIR}/")
    print(f"✓ Total applications analyzed: {metrics['total_applications']:,}")
    print(f"✓ Overall grant rate: {metrics['overall_grant_rate']:.1f}%")
    print(f"✓ Time span: {metrics['time_span']['start']}-{metrics['time_span']['end']}")

    return report

# Execute the analysis
if __name__ == "__main__":
    research_report = main()
